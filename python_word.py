from docx import Document
from docx.shared import Inches, Pt

#  步骤1.建立一个文档对象
doc = Document()  # 以默认模板建立文档对象

# doc = Document('a.docx')     # 读取a.docx文档，建立文档对象


#  步骤2.设置文档的格式（默认字体、页面边距等）
def chg_font(obj, fontname='微软雅黑', size=None):

    # 设置字体函数

    obj.font.name = fontname

    obj._element.rPr.rFonts.set(qn('w:eastAsia'), fontname)

    if size and isinstance(size, Pt):

        obj.font.size = size


distance = Inches(0.3)

sec = doc.sections[0]             # sections对应文档中的“节”

sec.left_margin = distance     # 以下依次设置左、右、上、下页面边距

sec.right_margin = distance

sec.top_margin = distance

sec.bottom_margin = distance

sec.page_width = Inches(12)  # 设置页面宽度

sec.page_height = Inches(20)  # 设置页面高度

# 设置默认字体

chg_font(doc.styles['Normal'], fontname='宋体')


# 步骤三.在文档对象中加入段落文本、表格、图像等，并指定其样式
# 1.添加段落文本
paragraph = doc.add_paragraph('text....')

ph_format = paragraph.paragraph_format

ph_format.space_before = Pt(10)  # 设置段前间距

ph_format.space_after = Pt(12)  # 设置段后间距

ph_format.line_spacing = Pt(19)  # 设置行间距
# 2.添加表格，并填写相关内容

tab = doc.add_table(rows=4, cols=4)  # 添加一个4行4列的空表

cell = tab.cell(1, 3)  # 获取某单元格对象（从0开始索引）
cell.text = 'abc'

# 步骤四.保存文档
doc.save('a.docx')
