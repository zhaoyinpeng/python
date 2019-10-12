import docx
import os
import sys

# 如果存在某文件就删除它
dirPath = ".//"
print('移除前目录下有文件：%s' % os.listdir(dirPath))
# 判断文件是否存在
if(os.path.exists(dirPath+"a.docx")):
    os.remove(dirPath+"a.docx")
    print('移除后test 目录下有文件：%s' % os.listdir(dirPath))
else:
    print('要删除的文件不存在！')
doc2 = docx.Document()  # 创建一个Document对象
doc2.add_paragraph('time')  # 增加一个paragraph
# 插入有序列表,段落的前面会有序号123
doc2.add_paragraph('把冰箱门打开', style='List Number')
doc2.add_paragraph('把大象装进去', style='List Number')
doc2.add_paragraph('把冰箱门关上', style='List Number')
# 插入无序列表，段落的前面没有序号
doc2.add_paragraph('把冰箱门打开2', style='List Bullet')
doc2.add_paragraph('把大象装进去2', style='List Bullet')
doc2.add_paragraph('把冰箱门关上2', style='List Bullet')
# 插入一个6行6列的表格
table = doc2.add_table(rows=6, cols=6, style='Table Grid')
for i in range(0, 6):
    for j in range(0, 6):
        table.cell(i, j).text = "第{i}行{j}列".format(i=i+1, j=j+1)

table2 = doc2.add_table(rows=6, cols=6, style='Table Grid')
for i in range(0, 6):
    for j in range(0, 6):
        table2.cell(i, j).text = "第{i}行{j}列".format(i=i+1, j=j+1)
# 插入照片
# doc2.add_picture('.\\girl.jpg', width=docx.shared.Inches(5))
print('保存文档为a.docx！')
doc2.save('a.docx')  # 保存文档

input()
